from django.shortcuts import render
from django.http import HttpResponse
from django.views import View
from django.urls import resolve
from django.core.paginator import Paginator
from django.contrib.auth.mixins import LoginRequiredMixin
from django.db.models import *
from django.db.models import Count,Sum
import json
import boto3
from django import forms
import docx

class details(forms.Form):
    name = forms.CharField(max_length=100,widget=forms.TextInput(attrs={'class':'input','placeholder':'Enter Name','name':'name'}))
    skills = forms.CharField(widget=forms.TextInput(attrs={'class':'input','placeholder':'Enter skills','name':'skills'}))

dynamodb = boto3.resource('dynamodb',aws_access_key_id='your_ak_here', aws_secret_access_key='your_sak_here', region_name='us-east-1')
s3 = boto3.resource('s3',aws_access_key_id='your_ak_here', aws_secret_access_key='your_sak_here', region_name='us-east-1')

class demo(View):
    def get(self, request, *args, **kwargs):
        return render(request, "showtemplates.html",{'demo':request})

class demo_submit(View):
    def get(self,request,*args,**kwargs):
        if(True):
            # table = dynamodb.Table('demo')
            # print(table)
            #client = boto3.client('s3',aws_access_key_id='', aws_secret_access_key='', region_name='us-east-1')
            #response = client.list_objects(Bucket='18suryateja.demo1')
            # for content in response['Contents']:
            #     obj_dict = client.get_object(Bucket='18suryateja.demo1', Key=content['Key'])
            #     print(content['Key'], obj_dict['LastModified'])
            #return HttpResponse(response['Key'])
            table = dynamodb.Table('demo')
            response = table.put_item(Item={'house':'white','name':'whstu'})
            path1 = "./template1.png"
            path2 = "./template2.png"
            s3.Bucket('18suryateja.demo1').upload_file(path1,Key='resumeTemplate1.png')
            s3.Bucket('18suryateja.demo1').upload_file(path2,Key='resumeTemplate2.png')
            return HttpResponse("21212")

class enter_details(View):
    def get(self,request,*args,**kwargs):
        return render(request,"details.html")
    def post(self,request,*args,**kwargs):
        print('post methodddd')
        return HttpResponse(kwargs['id'])

class get_resume(View):
    def post(self,request,*args,**kwargs):
        print('post methodddd')
        return HttpResponse(kwargs['id'])


class showTemplates(LoginRequiredMixin,View):
    login_url = '/login/'
    def get(self,request,*args,**kwargs):
        return render(request,"showtemplates.html")

class processSelectedTemplate(View):
    def get(self,request,*args,**kwargs):
        return render(request,'details.html',{'id':kwargs['id']})
    def post(self,request,*args,**kwargs):
        data = request.POST
        doc = docx.Document()
        if(str(kwargs['id'])=='1'):
            doc.add_heading(data['user_name'], 0)
            doc.add_heading('Email: '+ data['email'])
            doc.add_heading('Address: '+data['address'])
            doc.add_heading('Contact No.' + data['phone'])
            doc.add_heading('Skills')
            for skill in data['skills'].split(','):
                p0 = doc.add_paragraph(skill, style='List Bullet')
        # add a page break to start a new page
        #doc.add_page_break()
        # add a heading of level 2
            doc.add_heading("Workshops and Events:")
            doc_para= doc.add_paragraph(data['workshops'])
            doc.add_heading('Education')
            doc_para = doc.add_paragraph(' ')
            for ed in data['edu'].split(','):
                p0 = doc.add_paragraph(ed, style='List Bullet')
            doc.add_heading('Awards')
            doc_para = doc.add_paragraph(' ')
            for award in data['awards'].split(','):
                doc_para.add_run(award + ', ').bold = True
            doc.add_heading("Hobbies: ")
            doc_para = doc.add_paragraph(data['hobbies'])
            doc.add_heading("Languages Known:")
            doc_para = doc.add_paragraph(data['languages'])
        else:
            doc.add_heading(data['user_name'], 0)
            doc.add_heading('Email: ' + data['email'], 4)
            doc.add_heading('Contact No.' + data['phone'], 4)
            doc.add_heading('Education')
            table = doc.add_table(rows=4, cols=3)
            doc.add_paragraph(' ')
            num = 1
            table.cell(0,0).text = "Name of the College/School"
            table.cell(0,1).text = "Marks/CGPA"
            table.cell(0,2).text = "Specialization"
            table.cell(1, 0).text = data['edu'].split(',')[0]
            table.cell(2, 0).text = data['edu'].split(',')[1]
            table.cell(3,0).text =  data['edu'].split(',')[2]
            table.cell(1,1).text = "B.Tech"
            table.cell(2,1).text = "Board of Intermediate"
            table.cell(3,1).text = "10th Class"
            table.cell(1,2).text = '75';table.cell(2,2).text = "97"; table.cell(3,2).text = "85"
            table.style = 'Table Grid'
            doc.add_heading('Skills')
            for skill in data['skills'].split(','):
                p0 = doc.add_paragraph(skill, style='List Bullet')
        #doc.add_column(30)
        # pictures can also be added to our word document
        # width is optional
        #    doc.add_picture('path_to_picture')
        # now save the document to a location
            doc.add_heading("Technical Events:")
            doc_para = doc.add_paragraph(data['workshops'])
            doc.add_heading('Awards')
            doc_para = doc.add_paragraph(' ')
            for award in data['awards'].split(','):
                doc_para.add_run(award+', ').bold = True
            doc.add_heading('Address: ')
            doc_para = doc.add_paragraph(data['address'])
            doc.add_heading("Hobbies: ")
            doc_para = doc.add_paragraph(data['hobbies'])
            doc.add_heading("Languages Known:")
            doc_para = doc.add_paragraph(data['languages'])
        doc.save(data['name']+'.docx')
        res = s3.Bucket('resbucket587').upload_file(data['name']+'.docx', Key=data['name']+'.docx')
        table = dynamodb.Table('demo')
        response = table.put_item(Item={'house':'https://s3.ap-south-1.amazonaws.com/18suryateja.demo1/'+data['name'],'name':data['name']})
        bucket_file_path = "https://s3.ap-south-1.amazonaws.com/resbucket587/"+data['name']+'.docx'
        return HttpResponse(f"<br><br><center><a href={bucket_file_path}><button class='button'>Download Your Resume</button></a></center>")
